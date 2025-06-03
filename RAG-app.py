import os
from dotenv import load_dotenv
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain_openai import ChatOpenAI
from htmlTemplates import css, bot_template, user_template
from langchain.chains import ConversationalRetrievalChain

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    try:
        text=""
        pdf_reader=PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text+=page.extract_text()
        return text,len(pdf_reader.pages)
    except Exception as e:
        st.error(f"Error reading PDF {pdf_file.name}: {str(e)}")
        return "", 0




def extract_text_from_docx(docx_file):
    """Extract text from TXT file"""
    try:
        doc=Document(docx_file)
        text=""
        for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        return text, len(doc.paragraphs)
    except Exception as e:
        st.error(f"Error reading DOCX {docx_file.name}: {str(e)}")
        return "", 0
    


def extract_text_from_txt(txt_file):
    """Extract text from TXT file"""
    try:
        text=str(txt_file.read(),"utf-8")
        lines=text.count('\n')+1
        return text,lines
    except Exception as e:
        st.error(f"Error reading TXT {txt_file.name}:{str(e)}")
        return "",0



def get_documents_text(uploaded_files):
    """Extract text from multiple document types"""
    all_text = ""
    processed_files = []
    for file in uploaded_files:
        file_extension=file.name.split('.')[-1].lower()
        if file_extension == "pdf":
            text,pages = extract_text_from_pdf(file)
            file_info = f"PDF - {pages} pages"
        elif file_extension == 'docx':
            text, paragraphs = extract_text_from_docx(file)
            file_info = f"DOCX - {paragraphs} paragraphs"
        elif file_extension == 'txt':
            text, lines = extract_text_from_txt(file)
            file_info = f"TXT - {lines} lines"
        else:
            st.warning(f"Unsupported file type: {file.name}")
            continue

        if text.strip():
            all_text += f"\n\n--- From {file.name} ---\n"
            all_text += text
            processed_files.append({
                'name': file.name,
                'type': file_extension.upper(),
                'info': file_info,
                'status': 'Processed ✅'
            })

        else:
            processed_files.append({
                'name': file.name,
                'type': file_extension.upper(),
                'info': 'No text found',
                'status': 'Failed ❌'
            })
    return all_text, processed_files   
 


def get_text_chunks(text):
    text_splitter=CharacterTextSplitter(
        separator='\n',
        chunk_size=800,
        chunk_overlap=100,
        length_function=len
    )
    text_chunks=text_splitter.split_text(text)
    return text_chunks



def get_vectorstore(text_chunks):
    embeddings=OpenAIEmbeddings()
    vector_store=FAISS.from_texts(
        texts=text_chunks,
        embedding=embeddings
    )
    return vector_store


def get_conversation_chain(vectorstore):
    llm=ChatOpenAI(
        temperature=0.1,
        model_name="gpt-4o-mini",
        max_tokens=300
    )

    memory=ConversationBufferMemory(
        memory_key='chat_history', 
        return_messages=True,
        max_token_limit=1000
    )

    retriever=vectorstore.as_retriever(
        search_kwargs={"k": 3}
    )
    coversation_chain=ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        memory=memory,
        verbose=False
    )
    return coversation_chain
    

def handle_userinput(user_question):
    """Handle user input and generate response with system constraints"""
    if st.session_state.conversation is None:
        st.warning("Please upload and process documents first!")
        return
    
    enhanced_question = f"""
    SYSTEM INSTRUCTIONS:
    - Only answer questions based on the provided documents
    - If the answer is not in the documents, respond with "I don't know."
    - Keep responses concise and to the point
    - Do not make up or hallucinate information
    
    USER QUESTION: {user_question}
    """

    try:
        response = st.session_state.conversation({'question': enhanced_question})
        st.session_state.chat_history = response['chat_history']

        for i, message in enumerate(st.session_state.chat_history):
            if i % 2 == 0:
                user_msg=message.content
                actual_question = user_msg.split("USER QUESTION:")[-1].strip()
                st.write(user_template.replace(
                    "{{MSG}}",actual_question), unsafe_allow_html=True)
            else:
                st.write(bot_template.replace(
                    "{{MSG}}", message.content), unsafe_allow_html=True)
                
                
    except Exception as e:
        st.error(f"Error generating response: {str(e)}")
        st.info("💡 Tip: Try asking a more specific question about your documents.")

def reset_conversation():
    """ Reset chat history """
    st.session_state.conversation = None
    st.session_state.chat_history = None
    st.session_state.processed_files = []
    st.success("Successfully cleared chat history")



def main():
    load_dotenv()
    os.environ['OPENAI_API_KEY']=os.getenv('OPENAI_API_KEY')

    st.set_page_config(
        page_title="Chat with Multiple Documents",
        page_icon="📚",
        layout="wide"
    )

    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = None
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = []


    st.header("Chat with Multiple Documents 📚")
    st.markdown("*Upload PDF, DOCX, or TXT files and ask questions about their content*")

    user_question = st.text_input(
        "Ask a question about your documents:",
        placeholder="What is the main topic discussed in the documents?",
        key="user_input"
    )
    if user_question:
        handle_userinput(user_question)
   

    with st.sidebar:
        st.subheader("📁 Document Manager")
        
        # Multi-format file uploader
        uploaded_files = st.file_uploader(
            "Upload your documents here",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Supported formats: PDF, DOCX, TXT"
        )

        # List of uploaded files
        if uploaded_files:
            st.subheader("📋 Uploaded Files")
            for i, file in enumerate(uploaded_files):
                file_type = file.name.split('.')[-1].upper()
                file_size = f"{file.size / 1024:.1f} KB"
                st.write(f"**{i+1}.** {file.name}")
                st.write(f"   • Type: {file_type}")
                st.write(f"   • Size: {file_size}")
                st.write("---")
            
            st.write(f"**Total Files:** {len(uploaded_files)}")

        # Process button
        if st.button("🔄 Process Documents", type="primary"):
            if not uploaded_files:
                st.error("Please upload at least one document!")
            else:
                with st.spinner("Processing documents..."):
                    try:
                        # Get text from all documents
                        raw_text, processed_files = get_documents_text(uploaded_files)
                        
                        if raw_text.strip():
                            # Get text chunks
                            text_chunks = get_text_chunks(raw_text)
                            
                            # Create vector store
                            vectorstore = get_vectorstore(text_chunks)
                            
                            # Create conversation chain
                            st.session_state.conversation = get_conversation_chain(vectorstore)
                            st.session_state.processed_files = processed_files
                            
                            st.success(f"✅ Successfully processed {len([f for f in processed_files if 'Processed' in f['status']])} documents!")
                        else:
                            st.error("No text could be extracted from the uploaded files!")
                            
                    except Exception as e:
                        st.error(f"Error processing documents: {str(e)}")

        st.button('Reset Chat', on_click=reset_conversation)
    
    
    # Simple status messages
        if st.session_state.processed_files:
            st.subheader("📊 Processing Status")
            for file_info in st.session_state.processed_files:
                st.write(f"**{file_info['name']}**")
                st.write(f"   • Type: {file_info['type']}")
                st.write(f"   • Info: {file_info['info']}")
                st.write(f"   • Status: {file_info['status']}")
                st.write("---")
        
        # Instructions
        st.subheader("💡 How to Use")
        st.write("""
        1. **Upload** your documents (PDF, DOCX, TXT)
        2. **Review** the uploaded files list
        3. **Click** 'Process Documents' button
        4. **Ask** questions about your documents
        5. **Get** answers with source information
        """)

    

    

if __name__=='__main__':
    main()


